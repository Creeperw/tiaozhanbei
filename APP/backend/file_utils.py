import os
import json
from APP.backend.config import MAX_TEXT_LENGTH
from APP.backend.store import FILES

try:
    from docx import Document
except ImportError:
    Document = None

def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)

def read_text_file(file_path: str) -> str:
    encodings = ['utf-8', 'gbk', 'gb18030', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError: continue
    return "[无法识别文件编码]"

def extract_docx_content(file_path: str) -> str:
    if not Document: return "[Error: python-docx not installed]"
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip(): full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text: full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e: return f"[Error reading docx: {str(e)}]"

def get_file_content(file_id: str) -> str:
    """从磁盘读取文件并截断，用于 Prompt 拼接"""
    file_info = FILES.get(file_id)
    if not file_info: return "[文件未找到]"
    
    file_path = file_info.get("saved_path")
    if not file_path or not os.path.exists(file_path): return "[文件已从磁盘删除]"

    filename = file_info.get("original_name", "")
    ext = os.path.splitext(filename)[1].lower()
    content = ""

    try:
        if ext == '.docx': content = extract_docx_content(file_path)
        elif ext == '.json':
            raw = read_text_file(file_path)
            json_obj = json.loads(raw)
            content = json.dumps(json_obj, ensure_ascii=False, indent=2)
        else: content = read_text_file(file_path)
    except Exception as e: content = f"[读取出错: {str(e)}]"
        
    content = clean_text(content)
    if len(content) > MAX_TEXT_LENGTH:
        content = content[:MAX_TEXT_LENGTH] + f"\n\n... (截断: 剩余 {len(content) - MAX_TEXT_LENGTH} 字符)"
    
    return f"=== 文件内容: {filename} ===\n{content}\n=== 结束 ===\n"